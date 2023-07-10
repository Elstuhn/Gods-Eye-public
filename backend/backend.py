from flask import Flask, render_template, request, session, send_file, send_from_directory, jsonify
from flask_cors import CORS, cross_origin
import ai21
from docx import Document
from notion_client import Client
from pprint import pprint
from docx.shared import Inches
import requests
import redis
import json
import vertexai
from google.cloud import aiplatform
from vertexai.language_models import TextGenerationModel
from google.oauth2 import service_account

app = Flask(__name__)
cors = CORS(app)
app.config['CORS_HEADERS'] = 'Content-Type'

def get_transcription(url: str) -> str:
    value = r.get(url)
    if value is None:
        return ""
    else:
        return value.decode("utf-8")
    
def append_transcription(url: str, transcript: str):
    # if url in r: append to existing
    # else: create new
    existing = r.get(url)
    if existing is None:
        r.set(url, transcript)
    else:
        existing = existing.decode("utf-8")
        r.set(url, existing + ' ' + transcript)

def ask(transcript, prompt):
    parameters = {
    "temperature": 0.2,
    "max_output_tokens": 256,
    "top_p": 0.95,
    "top_k": 40
    }
    model = TextGenerationModel.from_pretrained("text-bison@001")
    response = model.predict(
        f"{prompt}: {transcript}",
        **parameters
    )
    return response.text
    

def write_text(client, page_id, text, type):
    client.blocks.children.append(
        block_id=page_id,
        children=[
            {
                "object": "block",
                "type": "paragraph",
                "paragraph": {
                    "rich_text": [
                        {
                            "type": "text",
                            "text": {
                                "content": text
                            }
                        }
                    ]
                }
            }
        ]
    )


@app.route('/notion', methods=["POST"])
def notion():
    url = request.get_json()['url']
    transcript = get_transcription(url)
    client = Client(auth=notion_token)

    write_text(client, notion_page_id, transcript, 'to_do')
    # Return a response
    return jsonify({"status": 1, "message": "Export to Notion completed successfully"})

def createDoc(transcript):
    summary = ask(transcript, "Provide a descriptive summary of this transcript")
    document = Document()
    document.add_heading('Meeting Minutes', 0)

    document.add_heading('Summary', level=1)
    document.add_paragraph(summary)
    document.add_paragraph("")
    answer = ask(transcript, "Give me a detailed summary of the meeting in bullet points")
    document.add_paragraph(answer)

    document.add_heading('Responsibilities', level=1)
    answer = ask(transcript, "What are the tasks and people responsible for them?")
    document.add_paragraph(
        answer
    )

    document.save('minutes.docx')
    return document

@app.route('/scribe', methods=["POST"])
def scribe():
    if request.method == "POST":
        path = "../minutes.docx"
        url = request.get_json()['url']
        transcript = get_transcription(url)
        createDoc(transcript)
        print('finished')
        return send_from_directory('../', 'minutes.docx', as_attachment=True)
    else:
        pass
    
@app.route('/clear', methods=["POST"])
def clear():
    if request.method == "POST":
        url = request.get_json()['url']
        r.delete(url)
        return ("OK", 200)
    else:
        pass

@app.route('/append-transcript', methods=["POST"])
def appendTranscript():
    if request.method == 'POST':
        body = request.get_json()
        if 'url' in body and 'stream' in body:
            append_transcription(body['url'], body['stream'])
            return ('OK', 200)
        else:
            return ('Bad Request', 400)

@app.route('/get-transcript', methods=["POST"])
def getTranscript():
    if request.method == 'POST':
        body = request.get_json()
        if 'url' in body:
            value = get_transcription(body['url'])
            return (value, 200)
        else:
            return ('Bad Request', 400)


if __name__ == '__main__':
    # conn_dict = psycopg.conninfo.conninfo_to_dict(connection)
    app.run(debug=True)
